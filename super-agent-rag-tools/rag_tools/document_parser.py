import base64
import hashlib
import html
import json
import re
from html.parser import HTMLParser
from io import BytesIO

from fastapi import HTTPException

from rag_tools.schemas.document_parse import (
    DocumentBlock,
    DocumentParseRequest,
    DocumentParseResponse,
    ParseArtifact,
    StructureNode,
    json_metadata,
)

PARSER_NAME = "rag-tools-python-parser"
PARSER_VERSION = "0.1.0"

NODE_TYPE_DOCUMENT = 1
NODE_TYPE_SECTION = 2


def parse_document(request: DocumentParseRequest) -> DocumentParseResponse:
    content = _decode_content(request.content_base64)
    file_type = _resolve_file_type(request.file_name, request.file_type, request.mime_type)

    if file_type in {"TXT", "MD"}:
        blocks = _parse_plain_text(content, markdown=file_type == "MD")
    elif file_type == "HTML":
        blocks = _parse_html(content)
    elif file_type == "PDF":
        blocks = _parse_pdf(content)
    elif file_type == "DOCX":
        blocks = _parse_docx(content)
    elif file_type == "XLSX":
        blocks = _parse_xlsx(content)
    elif file_type == "XLS":
        raise HTTPException(status_code=422, detail="XLS 解析未启用，请先转为 XLSX 后上传。")
    elif file_type == "DOC":
        raise HTTPException(status_code=422, detail="DOC 解析未启用，请先转为 DOCX 后上传。")
    else:
        raise HTTPException(status_code=422, detail=f"不支持的文件类型: {file_type or 'UNKNOWN'}")

    normalized_blocks = _normalize_blocks(blocks)
    _apply_table_context(normalized_blocks)
    parsed_text = _render_parsed_text(normalized_blocks)
    structure_nodes = _build_structure_nodes(request.file_name, normalized_blocks, parsed_text)
    paragraph_list = _paragraphs(parsed_text)
    heading_count = sum(1 for block in normalized_blocks if block.block_type == "TITLE")

    artifacts = _build_artifacts(request.file_name, parsed_text, normalized_blocks)

    return DocumentParseResponse(
        parsedText=parsed_text,
        charCount=len(parsed_text),
        tokenCount=_estimate_token_count(parsed_text),
        structureLevel=_structure_level(heading_count, len(paragraph_list)),
        contentQualityLevel=_content_quality_level(parsed_text),
        headingCount=heading_count,
        paragraphCount=len(paragraph_list),
        maxParagraphLength=max((len(item) for item in paragraph_list), default=0),
        artifacts=artifacts,
        blocks=normalized_blocks,
        structureNodes=structure_nodes,
    )


def _decode_content(content_base64: str) -> bytes:
    if not content_base64:
        raise HTTPException(status_code=422, detail="contentBase64 不能为空。")
    try:
        return base64.b64decode(content_base64, validate=True)
    except Exception as exception:
        raise HTTPException(status_code=422, detail="contentBase64 不是合法 Base64。") from exception


def _resolve_file_type(file_name: str, file_type: str, mime_type: str) -> str:
    explicit = (file_type or "").strip().upper()
    if explicit:
        return explicit
    suffix = (file_name or "").rsplit(".", 1)[-1].lower() if "." in (file_name or "") else ""
    if suffix == "pdf":
        return "PDF"
    if suffix == "docx":
        return "DOCX"
    if suffix == "doc":
        return "DOC"
    if suffix == "xlsx":
        return "XLSX"
    if suffix == "xls":
        return "XLS"
    if suffix in {"md", "markdown"}:
        return "MD"
    if suffix in {"html", "htm"}:
        return "HTML"
    if suffix == "txt" or (mime_type or "").startswith("text/"):
        return "TXT"
    return explicit


def _parse_plain_text(content: bytes, markdown: bool) -> list[DocumentBlock]:
    text = _decode_text(content)
    blocks = []
    block_no = 1
    for paragraph in _paragraphs(text):
        block_type = "TITLE" if markdown and paragraph.lstrip().startswith("#") else _classify_text_block(paragraph)
        clean_text = re.sub(r"^#{1,6}\s+", "", paragraph).strip()
        blocks.append(_block(block_no, block_type, clean_text))
        block_no += 1
    return blocks


def _parse_html(content: bytes) -> list[DocumentBlock]:
    parser = _HtmlBlockParser()
    parser.feed(_decode_text(content))
    return [
        _block(index + 1, block_type, text)
        for index, (block_type, text) in enumerate(parser.blocks)
        if text.strip()
    ]


def _parse_pdf(content: bytes) -> list[DocumentBlock]:
    try:
        import fitz
    except ImportError as exception:
        raise HTTPException(status_code=500, detail="Python parser 缺少 pymupdf 依赖，请执行 pip install -r requirements.txt。") from exception

    blocks = []
    with fitz.open(stream=content, filetype="pdf") as document:
        block_no = 1
        for page_index, page in enumerate(document, start=1):
            page_items: list[tuple[float, float, DocumentBlock]] = []
            table_regions = _extract_pdf_tables(page, page_index, block_no)
            for table_block in table_regions:
                bbox = _bbox_from_json(table_block.bbox_json)
                page_items.append((bbox[1], bbox[0], table_block))
                block_no += 1

            for raw_block in page.get_text("dict").get("blocks", []):
                bbox = tuple(float(value) for value in raw_block.get("bbox", (0, 0, 0, 0)))
                if raw_block.get("type") == 1:
                    image_block = _pdf_image_block(raw_block, block_no, page_index, bbox)
                    if image_block is not None:
                        page_items.append((bbox[1], bbox[0], image_block))
                        block_no += 1
                    continue
                if _bbox_inside_any(bbox, [_bbox_from_json(item.bbox_json) for item in table_regions]):
                    continue
                text = _cleanup_text(_pdf_text_block(raw_block))
                if not text:
                    continue
                page_items.append((
                    bbox[1],
                    bbox[0],
                    _block(block_no, _classify_pdf_text_block(text, raw_block), text,
                           page_no=page_index, bbox_json=_bbox_json(bbox),
                           metadata={"parser": PARSER_NAME, "layoutType": "text"}),
                ))
                block_no += 1

            page_items.sort(key=lambda item: (round(item[0], 1), round(item[1], 1)))
            blocks.extend(item[2] for item in page_items)
    return blocks


def _extract_pdf_tables(page, page_index: int, block_no: int) -> list[DocumentBlock]:
    if not hasattr(page, "find_tables"):
        return []
    try:
        tables = page.find_tables()
    except Exception:
        return []
    table_blocks: list[DocumentBlock] = []
    for table_index, table in enumerate(getattr(tables, "tables", []) or [], start=1):
        try:
            rows = _normalize_table_rows(table.extract())
        except Exception:
            rows = []
        if not _table_has_content(rows):
            continue
        bbox = tuple(float(value) for value in getattr(table, "bbox", (0, 0, 0, 0)))
        cell_metadata = _pdf_table_cell_metadata(table, rows, page_index)
        table_blocks.append(_block(
            block_no + len(table_blocks),
            "TABLE",
            _table_text(rows),
            page_no=page_index,
            bbox_json=_bbox_json(bbox),
            table_html=_table_html(rows),
            table_rows=rows,
            metadata={
                "parser": PARSER_NAME,
                "layoutType": "table",
                "tableNoOnPage": table_index,
                "engine": "pymupdf-find-tables",
                "tableCellMetadata": cell_metadata,
            },
        ))
    return table_blocks


def _pdf_text_block(raw_block: dict) -> str:
    lines = []
    for line in raw_block.get("lines", []):
        spans = []
        for span in line.get("spans", []):
            text = _cleanup_text(span.get("text", ""))
            if text:
                spans.append(text)
        if spans:
            lines.append("".join(spans))
    return "\n".join(lines)


def _classify_pdf_text_block(text: str, raw_block: dict) -> str:
    max_font_size = 0.0
    for line in raw_block.get("lines", []):
        for span in line.get("spans", []):
            max_font_size = max(max_font_size, float(span.get("size", 0.0) or 0.0))
    if len(text) <= 100 and max_font_size >= 14:
        return "TITLE"
    return _classify_text_block(text)


def _pdf_image_block(raw_block: dict, block_no: int, page_index: int, bbox: tuple[float, float, float, float]) -> DocumentBlock | None:
    image_bytes = raw_block.get("image")
    if not image_bytes:
        return None
    image_ext = str(raw_block.get("ext") or "png").lower()
    if image_ext not in {"png", "jpg", "jpeg", "webp"}:
        image_ext = "png"
    caption = f"第 {page_index} 页图片，位置 {round(bbox[0], 1)},{round(bbox[1], 1)}-{round(bbox[2], 1)},{round(bbox[3], 1)}。"
    return _block(
        block_no,
        "IMAGE",
        caption,
        page_no=page_index,
        bbox_json=_bbox_json(bbox),
        image_file_name=f"page-{page_index}-image-{block_no}.{image_ext}",
        image_content_base64=base64.b64encode(image_bytes).decode("ascii"),
        image_caption=caption,
        metadata={
            "parser": PARSER_NAME,
            "layoutType": "image",
            "imageExt": image_ext,
            "width": raw_block.get("width"),
            "height": raw_block.get("height"),
        },
    )


def _parse_docx(content: bytes) -> list[DocumentBlock]:
    try:
        from docx import Document
    except ImportError as exception:
        raise HTTPException(status_code=500, detail="Python parser 缺少 python-docx 依赖，请执行 pip install -r requirements.txt。") from exception

    document = Document(BytesIO(content))
    blocks = []
    block_no = 1
    for paragraph in document.paragraphs:
        text = _cleanup_text(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
        block_type = "TITLE" if "heading" in style_name.lower() else _classify_text_block(text)
        blocks.append(_block(block_no, block_type, text))
        block_no += 1

    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([_cleanup_text(cell.text) for cell in row.cells])
        if not rows:
            continue
        table_html = _table_html(rows)
        text = "\n".join(" | ".join(cell for cell in row if cell) for row in rows)
        blocks.append(_block(block_no, "TABLE", text, table_html=table_html, table_rows=rows))
        block_no += 1
    return blocks


def _parse_xlsx(content: bytes) -> list[DocumentBlock]:
    try:
        from openpyxl import load_workbook
    except ImportError as exception:
        raise HTTPException(status_code=500, detail="Python parser 缺少 openpyxl 依赖，请执行 pip install -r requirements.txt。") from exception

    workbook = load_workbook(BytesIO(content), read_only=False, data_only=True)
    blocks: list[DocumentBlock] = []
    block_no = 1
    for sheet in workbook.worksheets:
        merged_cells = _xlsx_merged_cell_lookup(sheet)
        extracted_rows = []
        for row_no, row in enumerate(sheet.iter_rows(values_only=False), start=1):
            extracted_row = []
            for column_no, cell in enumerate(row, start=1):
                source_row_no = _to_int(getattr(cell, "row", None)) or row_no
                source_column_no = _to_int(getattr(cell, "column", None)) or column_no
                merged_cell = merged_cells.get((source_row_no, source_column_no), {})
                extracted_row.append({
                    "value": _cell_text(getattr(cell, "value", None)),
                    "rowNo": source_row_no,
                    "columnNo": source_column_no,
                    "excelAddress": getattr(cell, "coordinate", "") or _excel_address(source_row_no, source_column_no),
                    "sheetName": sheet.title,
                    **merged_cell,
                })
            extracted_rows.append(extracted_row)
        rows, cell_metadata = _trim_table_rows_with_metadata(extracted_rows)
        if not _table_has_content(rows):
            continue
        rows, cell_metadata, table_title_rows, title_cell_metadata = _split_spreadsheet_title_rows(rows, cell_metadata)
        if not _table_has_content(rows):
            continue
        rows, cell_metadata, table_header_rows, header_cell_metadata = _flatten_spreadsheet_header_rows(rows, cell_metadata)
        rows, cell_metadata = _fill_spreadsheet_merged_data_values(rows, cell_metadata)
        title = f"工作表：{sheet.title}"
        blocks.append(_block(block_no, "TITLE", title, metadata={"parser": PARSER_NAME, "sheetName": sheet.title}))
        block_no += 1
        metadata = {
            "parser": PARSER_NAME,
            "layoutType": "spreadsheet",
            "sheetName": sheet.title,
            "maxRow": sheet.max_row,
            "maxColumn": sheet.max_column,
            "tableCellMetadata": cell_metadata,
        }
        if table_title_rows:
            metadata["tableTitleRows"] = table_title_rows
            metadata["tableTitleCellMetadata"] = title_cell_metadata
        if table_header_rows:
            metadata["tableHeaderFlattened"] = True
            metadata["tableHeaderRows"] = table_header_rows
            metadata["tableHeaderCellMetadata"] = header_cell_metadata
        blocks.append(_block(
            block_no,
            "TABLE",
            _table_text_with_titles(table_title_rows, rows),
            table_html=_table_html(rows),
            table_rows=rows,
            metadata=metadata,
        ))
        block_no += 1
    workbook.close()
    return blocks


def _normalize_blocks(blocks: list[DocumentBlock]) -> list[DocumentBlock]:
    normalized = []
    current_section = ""
    current_section_path = ""
    for index, block in enumerate(blocks, start=1):
        text = _cleanup_text(block.text)
        if not text and not block.table_html and not block.image_caption:
            continue
        block.block_no = len(normalized) + 1
        block.text = text
        if block.block_type == "TITLE":
            current_section = text
            current_section_path = text
        block.section_path = block.section_path or current_section_path
        block.canonical_path = block.canonical_path or _canonical_path(block.section_path, block.block_no)
        block.content_with_weight = block.content_with_weight or _content_with_weight(block, current_section)
        normalized.append(block)
    return normalized


def _apply_table_context(blocks: list[DocumentBlock]) -> None:
    for index, block in enumerate(blocks):
        if block.block_type != "TABLE":
            continue
        before = _nearby_context(blocks, index, -1)
        after = _nearby_context(blocks, index, 1)
        context_parts = []
        if block.section_path:
            context_parts.append(f"section: {block.section_path}")
        if before:
            context_parts.append(f"before: {before}")
        if after:
            context_parts.append(f"after: {after}")
        context_text = "\n".join(context_parts)
        if not context_text:
            continue
        _merge_block_metadata(block, {
            "tableContext": {
                "sectionPath": block.section_path,
                "before": before,
                "after": after,
            },
            "tableContextText": context_text,
        })
        if context_text not in block.content_with_weight:
            block.content_with_weight = "\n".join(part for part in [
                block.content_with_weight,
                "tableContext:",
                context_text,
            ] if part)


def _nearby_context(blocks: list[DocumentBlock], table_index: int, direction: int) -> str:
    items: list[str] = []
    index = table_index + direction
    while 0 <= index < len(blocks) and len(items) < 2:
        block = blocks[index]
        if block.block_type == "TABLE":
            break
        text = _cleanup_text(block.image_caption if block.block_type in {"IMAGE", "FIGURE"} else block.text)
        if text:
            items.append(_clip_text(text, 240))
        index += direction
    if direction < 0:
        items.reverse()
    return " / ".join(items)


def _merge_block_metadata(block: DocumentBlock, extra: dict) -> None:
    metadata = {}
    if block.metadata_json:
        try:
            metadata = json.loads(block.metadata_json)
            if not isinstance(metadata, dict):
                metadata = {"parserMetadata": metadata}
        except Exception:
            metadata = {"parserMetadata": block.metadata_json}
    metadata.update(extra)
    block.metadata_json = json_metadata(metadata)


def _clip_text(text: str, max_length: int) -> str:
    cleaned = _cleanup_text(text)
    if len(cleaned) <= max_length:
        return cleaned
    return cleaned[:max_length].rstrip() + "..."


def _build_structure_nodes(file_name: str, blocks: list[DocumentBlock], parsed_text: str) -> list[StructureNode]:
    nodes = [
        StructureNode(
            nodeNo=1,
            nodeType=NODE_TYPE_DOCUMENT,
            parentNodeNo=None,
            depth=0,
            nodeCode="ROOT",
            title=file_name or "document",
            anchorText=file_name or "document",
            canonicalPath="/",
            sectionPath="",
            contentText=parsed_text[:10000],
        )
    ]
    previous_heading_node_no = None
    for block in blocks:
        if block.block_type != "TITLE":
            continue
        node_no = len(nodes) + 1
        nodes.append(
            StructureNode(
                nodeNo=node_no,
                nodeType=NODE_TYPE_SECTION,
                parentNodeNo=1,
                prevSiblingNodeNo=previous_heading_node_no,
                depth=1,
                nodeCode=str(node_no - 1),
                title=block.text,
                anchorText=block.text[:200],
                canonicalPath=_canonical_path(block.text, node_no),
                sectionPath=block.text,
                contentText=block.text,
            )
        )
        if previous_heading_node_no is not None:
            nodes[-2].next_sibling_node_no = node_no
        previous_heading_node_no = node_no
    return nodes


def _build_artifacts(file_name: str, parsed_text: str, blocks: list[DocumentBlock]) -> list[ParseArtifact]:
    base_name = (file_name or "document").rsplit(".", 1)[0]
    markdown_bytes = parsed_text.encode("utf-8")
    blocks_json_bytes = json.dumps(
        [block.model_dump(by_alias=True) for block in blocks],
        ensure_ascii=False,
        indent=2,
    ).encode("utf-8")
    return [
        _artifact("MARKDOWN", f"{base_name}.md", "text/markdown;charset=UTF-8", markdown_bytes),
        _artifact("JSON", f"{base_name}.blocks.json", "application/json;charset=UTF-8", blocks_json_bytes),
    ]


def _artifact(artifact_type: str, file_name: str, content_type: str, content: bytes) -> ParseArtifact:
    return ParseArtifact(
        artifactType=artifact_type,
        fileName=file_name,
        contentType=content_type,
        contentBase64=base64.b64encode(content).decode("ascii"),
        contentHash=hashlib.sha256(content).hexdigest(),
        parserName=PARSER_NAME,
        parserVersion=PARSER_VERSION,
    )


def _block(
    block_no: int,
    block_type: str,
    text: str,
    *,
    page_no: int | None = None,
    bbox_json: str = "",
    table_html: str = "",
    table_rows: list[list[str]] | None = None,
    image_file_name: str = "",
    image_content_base64: str = "",
    image_caption: str = "",
    metadata: dict | None = None,
) -> DocumentBlock:
    return DocumentBlock(
        blockNo=block_no,
        blockType=block_type,
        pageNo=page_no,
        pageRange=str(page_no) if page_no is not None else "",
        bboxJson=bbox_json,
        text=text,
        tableHtml=table_html,
        tableRows=table_rows or [],
        imageFileName=image_file_name,
        imageContentBase64=image_content_base64,
        imageCaption=image_caption,
        metadataJson=json_metadata(metadata or {"parser": PARSER_NAME}),
    )


def _render_parsed_text(blocks: list[DocumentBlock]) -> str:
    parts = []
    for block in blocks:
        if block.block_type == "TITLE":
            parts.append(f"# {block.text}")
        elif block.block_type == "TABLE" and block.table_html:
            parts.append(block.text)
        elif block.block_type in {"IMAGE", "FIGURE"} and block.image_caption:
            parts.append(block.image_caption)
        else:
            parts.append(block.text)
    return _cleanup_text("\n\n".join(part for part in parts if part))


def _content_with_weight(block: DocumentBlock, current_section: str) -> str:
    parts = []
    if current_section:
        parts.append(f"section: {current_section}")
    if block.block_type:
        parts.append(f"type: {block.block_type}")
    if block.text:
        parts.append(block.text)
    if block.image_caption and block.image_caption != block.text:
        parts.append(f"caption: {block.image_caption}")
    if block.table_rows:
        parts.append(_table_text(block.table_rows))
    return "\n".join(parts)


def _canonical_path(section_path: str, block_no: int) -> str:
    section = re.sub(r"[^0-9a-zA-Z\u4e00-\u9fff]+", "-", section_path or "root").strip("-")
    return f"/{section or 'root'}/{block_no}"


def _bbox_json(bbox: tuple[float, float, float, float]) -> str:
    return json_metadata({
        "x0": round(float(bbox[0]), 2),
        "y0": round(float(bbox[1]), 2),
        "x1": round(float(bbox[2]), 2),
        "y1": round(float(bbox[3]), 2),
    })


def _bbox_from_json(bbox_json: str) -> tuple[float, float, float, float]:
    if not bbox_json:
        return 0.0, 0.0, 0.0, 0.0
    try:
        data = json.loads(bbox_json)
        return (
            float(data.get("x0", 0.0)),
            float(data.get("y0", 0.0)),
            float(data.get("x1", 0.0)),
            float(data.get("y1", 0.0)),
        )
    except Exception:
        return 0.0, 0.0, 0.0, 0.0


def _bbox_inside_any(bbox: tuple[float, float, float, float], regions: list[tuple[float, float, float, float]]) -> bool:
    x0, y0, x1, y1 = bbox
    for rx0, ry0, rx1, ry1 in regions:
        overlap_x = max(0.0, min(x1, rx1) - max(x0, rx0))
        overlap_y = max(0.0, min(y1, ry1) - max(y0, ry0))
        overlap_area = overlap_x * overlap_y
        area = max(1.0, (x1 - x0) * (y1 - y0))
        if overlap_area / area >= 0.65:
            return True
    return False


def _normalize_table_rows(rows) -> list[list[str]]:
    return _trim_table_rows([
        [_cell_text(cell) for cell in row]
        for row in (rows or [])
    ])


def _pdf_table_cell_metadata(table, rows: list[list[str]], page_index: int) -> list[dict]:
    cells = _pdf_table_cells(table)
    column_count = max((len(row) for row in rows), default=0)
    metadata: list[dict] = []
    for row_index, row in enumerate(rows):
        for column_index, cell_text in enumerate(row):
            item = {
                "rowNo": row_index + 1,
                "columnNo": column_index + 1,
                "sourceRowNo": row_index + 1,
                "sourceColumnNo": column_index + 1,
                "cellCoordinate": f"R{row_index + 1}C{column_index + 1}",
                "pageNo": page_index,
                "value": cell_text,
            }
            bbox = _pdf_table_cell_bbox(cells, row_index, column_index, column_count)
            if bbox is not None:
                item["bboxJson"] = _bbox_json(bbox)
            metadata.append(item)
    return metadata


def _pdf_table_cells(table) -> list:
    cells = getattr(table, "cells", None)
    if not cells:
        return []
    if isinstance(cells, list):
        return cells
    try:
        return list(cells)
    except TypeError:
        return []


def _pdf_table_cell_bbox(cells: list,
                         row_index: int,
                         column_index: int,
                         column_count: int) -> tuple[float, float, float, float] | None:
    for cell in cells:
        row_no = _first_int_attr(cell, "row", "row_no", "rowno", "r")
        column_no = _first_int_attr(cell, "col", "column", "column_no", "colno", "c")
        if row_no is None or column_no is None:
            continue
        if _zero_based(row_no, row_index) and _zero_based(column_no, column_index):
            bbox = _coerce_bbox(getattr(cell, "bbox", cell))
            if bbox is not None:
                return bbox
    sorted_bboxes = _sorted_cell_bboxes(cells)
    if sorted_bboxes:
        linear_index = row_index * max(column_count, 1) + column_index
        if 0 <= linear_index < len(sorted_bboxes):
            return sorted_bboxes[linear_index]
    linear_index = row_index * max(column_count, 1) + column_index
    if 0 <= linear_index < len(cells):
        return _coerce_bbox(getattr(cells[linear_index], "bbox", cells[linear_index]))
    return None


def _sorted_cell_bboxes(cells: list) -> list[tuple[float, float, float, float]]:
    bboxes = []
    for cell in cells:
        bbox = _coerce_bbox(getattr(cell, "bbox", cell))
        if bbox is not None:
            bboxes.append(bbox)
    return sorted(bboxes, key=lambda bbox: (round(bbox[1], 1), round(bbox[0], 1)))


def _first_int_attr(value, *names: str) -> int | None:
    if isinstance(value, dict):
        for name in names:
            parsed = _to_int(value.get(name))
            if parsed is not None:
                return parsed
        return None
    for name in names:
        parsed = _to_int(getattr(value, name, None))
        if parsed is not None:
            return parsed
    return None


def _to_int(value) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _zero_based(value: int, expected_zero_based: int) -> bool:
    return value == expected_zero_based or value == expected_zero_based + 1


def _coerce_bbox(value) -> tuple[float, float, float, float] | None:
    if value is None:
        return None
    if isinstance(value, dict):
        keys = ("x0", "y0", "x1", "y1")
        if all(key in value for key in keys):
            try:
                return tuple(float(value[key]) for key in keys)
            except (TypeError, ValueError):
                return None
        return None
    if hasattr(value, "x0") and hasattr(value, "y0") and hasattr(value, "x1") and hasattr(value, "y1"):
        try:
            return float(value.x0), float(value.y0), float(value.x1), float(value.y1)
        except (TypeError, ValueError):
            return None
    if isinstance(value, (list, tuple)) and len(value) >= 4:
        try:
            return float(value[0]), float(value[1]), float(value[2]), float(value[3])
        except (TypeError, ValueError):
            return None
    return None


def _trim_table_rows_with_metadata(rows: list[list[dict]]) -> tuple[list[list[str]], list[dict]]:
    cleaned = [
        [
            {
                **cell,
                "value": _cleanup_text(str(cell.get("value") or "")),
            }
            for cell in row
        ]
        for row in rows
    ]
    cleaned = [row for row in cleaned if any(cell.get("value") for cell in row)]
    if not cleaned:
        return [], []
    max_columns = max(len(row) for row in cleaned)
    padded = [row + [_empty_cell_metadata() for _ in range(max_columns - len(row))] for row in cleaned]
    non_empty_columns = [
        index
        for index in range(max_columns)
        if any(row[index].get("value") for row in padded)
    ]
    if not non_empty_columns:
        return [], []
    first_column = min(non_empty_columns)
    last_column = max(non_empty_columns)
    trimmed_rows = [row[first_column:last_column + 1] for row in padded]
    values = [[cell.get("value", "") for cell in row] for row in trimmed_rows]
    metadata = []
    for row_index, row in enumerate(trimmed_rows, start=1):
        for column_index, cell in enumerate(row, start=1):
            source_row_no = _to_int(cell.get("rowNo"))
            source_column_no = _to_int(cell.get("columnNo"))
            item = {
                "rowNo": row_index,
                "columnNo": column_index,
                "sourceRowNo": source_row_no,
                "sourceColumnNo": source_column_no,
                "excelAddress": cell.get("excelAddress") or "",
                "cellCoordinate": cell.get("excelAddress") or f"R{source_row_no or row_index}C{source_column_no or column_index}",
                "sheetName": cell.get("sheetName") or "",
                "value": cell.get("value", ""),
            }
            for key in (
                "mergedCell",
                "mergedCellAnchor",
                "mergedCellRange",
                "mergedCellAnchorAddress",
                "mergedCellValue",
                "rowSpan",
                "columnSpan",
            ):
                if key in cell:
                    item[key] = cell.get(key)
            metadata.append({key: value for key, value in item.items() if value not in (None, "")})
    return values, metadata


def _xlsx_merged_cell_lookup(sheet) -> dict[tuple[int, int], dict]:
    merged_cells: dict[tuple[int, int], dict] = {}
    merged_ranges = getattr(getattr(sheet, "merged_cells", None), "ranges", []) or []
    for merged_range in merged_ranges:
        min_row = int(getattr(merged_range, "min_row", 0) or 0)
        max_row = int(getattr(merged_range, "max_row", 0) or 0)
        min_col = int(getattr(merged_range, "min_col", 0) or 0)
        max_col = int(getattr(merged_range, "max_col", 0) or 0)
        if min_row <= 0 or min_col <= 0 or max_row < min_row or max_col < min_col:
            continue
        anchor_address = _excel_address(min_row, min_col)
        anchor_value = _cell_text(sheet.cell(row=min_row, column=min_col).value)
        row_span = max_row - min_row + 1
        column_span = max_col - min_col + 1
        for row_no in range(min_row, max_row + 1):
            for column_no in range(min_col, max_col + 1):
                is_anchor = row_no == min_row and column_no == min_col
                merged_cells[(row_no, column_no)] = {
                    "mergedCell": True,
                    "mergedCellAnchor": is_anchor,
                    "mergedCellRange": str(getattr(merged_range, "coord", "")),
                    "mergedCellAnchorAddress": anchor_address,
                    "mergedCellValue": anchor_value,
                    "rowSpan": row_span if is_anchor else 1,
                    "columnSpan": column_span if is_anchor else 1,
                }
    return merged_cells


def _split_spreadsheet_title_rows(rows: list[list[str]], metadata: list[dict]) -> tuple[list[list[str]], list[dict], list[str], list[dict]]:
    title_rows: list[str] = []
    title_metadata: list[dict] = []
    remaining_rows = rows
    remaining_metadata = metadata
    while len(remaining_rows) > 1 and _is_spreadsheet_title_row(remaining_rows[0], remaining_rows[1], remaining_metadata):
        title_rows.append(" | ".join(cell for cell in remaining_rows[0] if cell).strip())
        title_metadata.extend([item for item in remaining_metadata if item.get("rowNo") == 1])
        remaining_rows = remaining_rows[1:]
        remaining_metadata = _shift_table_cell_metadata_rows(remaining_metadata, 1)
    return remaining_rows, remaining_metadata, title_rows, title_metadata


def _is_spreadsheet_title_row(row: list[str], next_row: list[str], metadata: list[dict]) -> bool:
    non_blank_cells = [cell for cell in row if _cleanup_text(cell)]
    next_non_blank_count = sum(1 for cell in next_row if _cleanup_text(cell))
    if len(non_blank_cells) != 1 or next_non_blank_count < 2:
        return False
    first_row_metadata = [item for item in metadata if item.get("rowNo") == 1]
    if not first_row_metadata:
        return False
    anchor = next((item for item in first_row_metadata if item.get("value") == non_blank_cells[0]), first_row_metadata[0])
    if not anchor.get("mergedCell"):
        return False
    return int(anchor.get("columnSpan") or 1) >= min(2, len(row))


def _shift_table_cell_metadata_rows(metadata: list[dict], removed_rows: int) -> list[dict]:
    shifted = []
    for item in metadata:
        row_no = _to_int(item.get("rowNo"))
        if row_no is None or row_no <= removed_rows:
            continue
        shifted_item = {**item, "rowNo": row_no - removed_rows}
        shifted.append(shifted_item)
    return shifted


def _flatten_spreadsheet_header_rows(rows: list[list[str]], metadata: list[dict]) -> tuple[list[list[str]], list[dict], list[list[str]], list[dict]]:
    if not _is_spreadsheet_multi_level_header(rows, metadata):
        return rows, metadata, [], []
    column_count = max((len(row) for row in rows[:2]), default=0)
    flattened_header: list[str] = []
    flattened_metadata: list[dict] = []
    metadata_by_position = _metadata_by_position(metadata)
    for column_index in range(column_count):
        column_no = column_index + 1
        top_value = _effective_header_value(rows, metadata_by_position, 1, column_no)
        child_value = _effective_header_value(rows, metadata_by_position, 2, column_no)
        header_name = _join_header_parts(top_value, child_value) or f"列{column_no}"
        source_metadata = _header_source_metadata(metadata_by_position, rows, column_no)
        flattened_metadata.append({
            **source_metadata,
            "rowNo": 1,
            "columnNo": column_no,
            "value": header_name,
            "flattenedHeader": True,
            "headerSourceRows": 2,
        })
        flattened_header.append(header_name)
    shifted_data_metadata = []
    for item in metadata:
        row_no = _to_int(item.get("rowNo"))
        if row_no is None or row_no <= 2:
            continue
        shifted_data_metadata.append({**item, "rowNo": row_no - 1})
    return [flattened_header] + rows[2:], flattened_metadata + shifted_data_metadata, rows[:2], [
        item for item in metadata if _to_int(item.get("rowNo")) in {1, 2}
    ]


def _fill_spreadsheet_merged_data_values(rows: list[list[str]], metadata: list[dict]) -> tuple[list[list[str]], list[dict]]:
    if len(rows) < 2:
        return rows, metadata
    filled_rows = [list(row) for row in rows]
    metadata_by_position = _metadata_by_position(metadata)
    filled_positions: set[tuple[int, int]] = set()
    for row_index in range(1, len(filled_rows)):
        row_no = row_index + 1
        for column_index, cell_text in enumerate(filled_rows[row_index]):
            column_no = column_index + 1
            if _cleanup_text(cell_text):
                continue
            cell_metadata = metadata_by_position.get((row_no, column_no), {})
            if not cell_metadata.get("mergedCell"):
                continue
            merged_value = _cleanup_text(str(cell_metadata.get("mergedCellValue") or ""))
            if not merged_value:
                continue
            filled_rows[row_index][column_index] = merged_value
            filled_positions.add((row_no, column_no))
    if not filled_positions:
        return rows, metadata
    filled_metadata = []
    for item in metadata:
        row_no = _to_int(item.get("rowNo"))
        column_no = _to_int(item.get("columnNo"))
        if (row_no, column_no) in filled_positions:
            item = {
                **item,
                "value": filled_rows[row_no - 1][column_no - 1],
                "mergedValueFilled": True,
            }
        filled_metadata.append(item)
    return filled_rows, filled_metadata


def _is_spreadsheet_multi_level_header(rows: list[list[str]], metadata: list[dict]) -> bool:
    if len(rows) < 3:
        return False
    column_count = max((len(row) for row in rows[:2]), default=0)
    if column_count < 2:
        return False
    first_row_metadata = [item for item in metadata if item.get("rowNo") == 1]
    has_group_header = any(
        item.get("mergedCell")
        and _to_int(item.get("columnSpan")) is not None
        and _to_int(item.get("columnSpan")) > 1
        and _not_blank(item.get("mergedCellValue") or item.get("value"))
        for item in first_row_metadata
    )
    if not has_group_header:
        return False
    return _non_blank_cell_count(rows[1]) >= 2 and _non_blank_cell_count(rows[2]) >= 2


def _metadata_by_position(metadata: list[dict]) -> dict[tuple[int, int], dict]:
    result = {}
    for item in metadata:
        row_no = _to_int(item.get("rowNo"))
        column_no = _to_int(item.get("columnNo"))
        if row_no is not None and column_no is not None:
            result[(row_no, column_no)] = item
    return result


def _effective_header_value(rows: list[list[str]], metadata_by_position: dict[tuple[int, int], dict], row_no: int, column_no: int) -> str:
    row_index = row_no - 1
    column_index = column_no - 1
    value = ""
    if 0 <= row_index < len(rows) and 0 <= column_index < len(rows[row_index]):
        value = rows[row_index][column_index]
    value = _cleanup_text(value)
    if value:
        return value
    metadata = metadata_by_position.get((row_no, column_no), {})
    return _cleanup_text(str(metadata.get("mergedCellValue") or ""))


def _header_source_metadata(metadata_by_position: dict[tuple[int, int], dict], rows: list[list[str]], column_no: int) -> dict:
    child_value = _effective_header_value(rows, metadata_by_position, 2, column_no)
    top_value = _effective_header_value(rows, metadata_by_position, 1, column_no)
    source = metadata_by_position.get((2, column_no), {}) if child_value else metadata_by_position.get((1, column_no), {})
    if not source and top_value:
        source = metadata_by_position.get((1, column_no), {})
    allowed = {
        "sourceRowNo",
        "sourceColumnNo",
        "excelAddress",
        "cellCoordinate",
        "sheetName",
        "mergedCell",
        "mergedCellAnchor",
        "mergedCellRange",
        "mergedCellAnchorAddress",
        "mergedCellValue",
        "rowSpan",
        "columnSpan",
    }
    return {key: value for key, value in source.items() if key in allowed}


def _join_header_parts(*parts: str) -> str:
    result = []
    for part in parts:
        cleaned = _cleanup_text(part)
        if cleaned and cleaned not in result:
            result.append(cleaned)
    return " ".join(result)


def _non_blank_cell_count(row: list[str]) -> int:
    return sum(1 for cell in row if _cleanup_text(cell))


def _not_blank(value) -> bool:
    return bool(_cleanup_text(str(value or "")))


def _empty_cell_metadata() -> dict:
    return {"value": ""}


def _trim_table_rows(rows: list[list[str]]) -> list[list[str]]:
    cleaned = [[_cleanup_text(cell) for cell in row] for row in rows]
    cleaned = [row for row in cleaned if any(cell for cell in row)]
    if not cleaned:
        return []
    max_columns = max(len(row) for row in cleaned)
    padded = [row + [""] * (max_columns - len(row)) for row in cleaned]
    non_empty_columns = [
        index
        for index in range(max_columns)
        if any(row[index] for row in padded)
    ]
    if not non_empty_columns:
        return []
    first_column = min(non_empty_columns)
    last_column = max(non_empty_columns)
    return [row[first_column:last_column + 1] for row in padded]


def _table_has_content(rows: list[list[str]]) -> bool:
    return any(any(cell.strip() for cell in row) for row in rows or [])


def _table_text(rows: list[list[str]]) -> str:
    return "\n".join(" | ".join(cell for cell in row if cell) for row in rows if any(row))


def _table_text_with_titles(title_rows: list[str], rows: list[list[str]]) -> str:
    parts = [title for title in title_rows if title]
    table_text = _table_text(rows)
    if table_text:
        parts.append(table_text)
    return "\n".join(parts)


def _cell_text(value) -> str:
    if value is None:
        return ""
    return _cleanup_text(str(value))


def _excel_address(row_no: int, column_no: int) -> str:
    return f"{_excel_column_name(column_no)}{row_no}"


def _excel_column_name(column_no: int) -> str:
    if column_no <= 0:
        return ""
    name = ""
    while column_no > 0:
        column_no, remainder = divmod(column_no - 1, 26)
        name = chr(65 + remainder) + name
    return name


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _cleanup_text(text: str) -> str:
    if not text:
        return ""
    return (
        text.replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\x00", " ")
        .replace("\t", " ")
        .strip()
    )


def _paragraphs(text: str) -> list[str]:
    cleaned = _cleanup_text(text)
    if not cleaned:
        return []
    return [item.strip() for item in re.split(r"\n\s*\n+", cleaned) if item.strip()]


def _classify_text_block(text: str) -> str:
    stripped = text.strip()
    if not stripped:
        return "TEXT"
    if len(stripped) <= 80 and (
        re.match(r"^([一二三四五六七八九十]+、|第[一二三四五六七八九十0-9]+[章节条]|[0-9]+(\.[0-9]+)*[、.])", stripped)
        or stripped.endswith(("章", "节", "：", ":"))
    ):
        return "TITLE"
    return "TEXT"


def _estimate_token_count(text: str) -> int:
    if not text:
        return 0
    chinese_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    english_count = len(re.findall(r"[A-Za-z]+", text))
    return chinese_count + english_count + max(1, (len(text) - chinese_count) // 4)


def _structure_level(heading_count: int, paragraph_count: int) -> int:
    if heading_count >= 5:
        return 3
    if heading_count >= 2:
        return 2
    if paragraph_count >= 3:
        return 1
    return 0


def _content_quality_level(text: str) -> int:
    if not text or len(text) < 20:
        return 1
    broken_ratio = text.count("�") / max(len(text), 1)
    if broken_ratio > 0.02 or len(text) < 100:
        return 1
    if broken_ratio > 0.005 or len(text) < 500:
        return 2
    return 3


def _table_html(rows: list[list[str]]) -> str:
    body = []
    for row in rows:
        cells = "".join(f"<td>{html.escape(cell)}</td>" for cell in row)
        body.append(f"<tr>{cells}</tr>")
    return "<table><tbody>" + "".join(body) + "</tbody></table>"


class _HtmlBlockParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[tuple[str, str]] = []
        self._current_tag = ""
        self._buffer: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li"}:
            self._flush()
            self._current_tag = tag

    def handle_endtag(self, tag: str) -> None:
        if tag == self._current_tag:
            self._flush()

    def handle_data(self, data: str) -> None:
        if self._current_tag:
            self._buffer.append(data)

    def _flush(self) -> None:
        text = _cleanup_text("".join(self._buffer))
        if text:
            block_type = "TITLE" if self._current_tag.startswith("h") else _classify_text_block(text)
            self.blocks.append((block_type, text))
        self._buffer = []
        self._current_tag = ""
